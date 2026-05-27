"""Tests for context builder media handling.

The ContextBuilder._build_user_content method should ONLY handle images.
Document text extraction is the responsibility of the processing layer
(AgentLoop._process_message and _drain_pending).
"""

from __future__ import annotations

from pathlib import Path

from nanobot.agent.context import ContextBuilder
from nanobot.utils.document import extract_documents


def _make_builder(tmp_path: Path) -> ContextBuilder:
    """Create a minimal ContextBuilder for testing."""
    return ContextBuilder(workspace=tmp_path, timezone="UTC")


def test_build_user_content_with_no_media_returns_string(tmp_path: Path) -> None:
    builder = _make_builder(tmp_path)
    result = builder._build_user_content("hello", None)
    assert result == "hello"


def test_build_user_content_with_image_returns_list(tmp_path: Path) -> None:
    """Image files should produce base64 content blocks."""
    builder = _make_builder(tmp_path)
    png = tmp_path / "test.png"
    png.write_bytes(b"\x89PNG\r\n\x1a\n" + b"\x00" * 100)
    result = builder._build_user_content("describe this", [str(png)])
    assert isinstance(result, list)
    types = [b["type"] for b in result]
    assert "image_url" in types
    assert "text" in types


def test_build_user_content_ignores_non_image_files(tmp_path: Path) -> None:
    """Non-image files should be silently skipped — extraction is not context builder's job."""
    builder = _make_builder(tmp_path)
    txt = tmp_path / "notes.txt"
    txt.write_text("some text", encoding="utf-8")
    result = builder._build_user_content("summarize", [str(txt)])
    assert result == "summarize"


def test_build_user_content_mixed_image_and_non_image(tmp_path: Path) -> None:
    """Only images should be included; non-image files are skipped."""
    builder = _make_builder(tmp_path)
    png = tmp_path / "chart.png"
    png.write_bytes(b"\x89PNG\r\n\x1a\n" + b"\x00" * 100)
    txt = tmp_path / "report.txt"
    txt.write_text("report text", encoding="utf-8")

    result = builder._build_user_content("analyze", [str(png), str(txt)])
    assert isinstance(result, list)
    assert any(b["type"] == "image_url" for b in result)
    text_parts = [b.get("text", "") for b in result if b.get("type") == "text"]
    assert all("report text" not in t for t in text_parts)


# ---------------------------------------------------------------------------
# Bug detection: extract_documents must be called BEFORE _build_user_content
# to prevent document media from being silently dropped.
# This simulates the _drain_pending code path.
# ---------------------------------------------------------------------------

def test_drain_pending_path_preserves_document_text(tmp_path: Path) -> None:
    """Simulates the _drain_pending path: a pending follow-up message
    with a document attachment must have its text extracted before being
    passed to _build_user_content.  Without extract_documents, the
    document is silently dropped."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Quarterly revenue is $5M")
    docx_path = tmp_path / "report.docx"
    doc.save(docx_path)

    content = "summarize"
    media = [str(docx_path)]

    # Step 1: extract_documents separates docs from images
    new_content, image_only = extract_documents(content, media)

    # Step 2: _build_user_content handles only images (none left here)
    builder = _make_builder(tmp_path)
    result = builder._build_user_content(new_content, image_only if image_only else None)

    # The document text should be present in the final content
    assert "Quarterly revenue" in result
    assert "summarize" in result


def test_drain_pending_path_without_extract_loses_document(tmp_path: Path) -> None:
    """Demonstrates the BUG: if _drain_pending calls _build_user_content
    directly without extract_documents, document content is lost."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Secret data in document")
    docx_path = tmp_path / "report.docx"
    doc.save(docx_path)

    builder = _make_builder(tmp_path)

    # Bug path: call _build_user_content directly with document media
    result = builder._build_user_content("summarize", [str(docx_path)])

    # The document text is LOST — _build_user_content ignores non-images
    assert result == "summarize"  # only the original text, no doc content
    assert "Secret data" not in result
